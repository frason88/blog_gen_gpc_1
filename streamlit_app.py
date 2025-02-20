import streamlit as st
import os
import re
import shutil
import replicate
import random
import time
import backoff
import json
import hashlib
import tempfile
from docx import Document
from docx.shared import Inches
import markdown
import io
import re
import requests
from bs4 import BeautifulSoup

def sanitize_filename(filename):
    filename = re.sub(r'[<>:"/\\|?*]', '_', filename)
    if len(filename) > 100:
        filename = hashlib.md5(filename.encode()).hexdigest()
    return filename


def introduce_grammatical_errors(text, error_rate=0.05):
    sentences = re.split(r'(?<!\w\.\w.)(?<![A-Z][a-z]\.)(?<=\.|\?)\s', text)
    modified_sentences = []

    for sentence in sentences:
        if random.random() < error_rate:
            words = sentence.split()
            if len(words) > 2:
                if random.random() < 0.5:
                    if words[0].lower() not in ['i', 'we', 'you', 'they'] and not words[0].endswith('s'):
                        if words[1].lower() == 'is':
                            words[1] = 'are'
                        elif words[1].lower() == 'was':
                            words[1] = 'were'
                    sentence = ' '.join(words)
                else:
                    if len(words) > 5:
                        split_index = random.randint(1, len(words) - 2)
                        words.insert(split_index, ',')
                        sentence = ' '.join(words)
        modified_sentences.append(sentence)

    return ' '.join(modified_sentences)


def generate_image(prompt, image_path, api_key):
    try:
        # Set the API key for the current session
        os.environ["REPLICATE_API_TOKEN"] = api_key

        image_url = replicate.run(
            "google/imagen-3-fast",
            input={
                "prompt": prompt,
                "aspect_ratio": "16:9",
                "safety_filter_level": "block_medium_and_above"
            }
        )
        if image_url:
            print(f"Image generated successfully. URL: {image_url}")
        else:
            print("Image generation failed.")
            return None
        return image_url
    except Exception as e:
        print(f"Error generating image: {e}")
        return None


@backoff.on_exception(backoff.expo, Exception, max_tries=2, max_time=300)
def run_with_retry(runner, topic_with_keywords):
    try:
        runner.run(
            topic=topic_with_keywords,
            do_research=True,
            do_generate_outline=True,
            do_generate_article=True,
            do_polish_article=False,
        )
        runner.post_run()
    except Exception as e:
        print(f"Error during STORM process: {e}. Waiting 10 seconds before retrying...")
        time.sleep(60)
        raise


def format_references(references_data):
    if not isinstance(references_data, dict):
        return "References data is not in the expected format."

    # Try to access url_to_info if it exists
    if "url_to_info" in references_data:
        references_dict = references_data["url_to_info"]
    else:
        references_dict = references_data

    # Create a numbered list of references with each reference on its own line
    formatted_references = []
    for i, (url, data) in enumerate(references_dict.items(), 1):
        if isinstance(data, dict) and 'title' in data and 'url' in data:
            # Format with number that matches the citation in text and make it a link
            formatted_references.append(f"[{i}] [{data['title']}]({data['url']})")
        else:
            formatted_references.append(f"[{i}] Reference information missing for {url}")

    # Join with double newlines to ensure one citation per line with spacing
    return "\n\n".join(formatted_references)


def process_citations(article_text):
    """Convert citation numbers to linked citations"""
    import re

    # This regex finds citation numbers in square brackets
    citation_pattern = r'\[(\d+)\]'

    # Function to replace each citation match with a linked version
    def replace_citation(match):
        citation_num = match.group(1)
        return f'[{citation_num}](#references)'

    # Replace all citations with linked versions
    processed_text = re.sub(citation_pattern, replace_citation, article_text)

    return processed_text


def main_in_memory(topic, keywords, anthropic_api_key, serper_api_key, replicate_api_key):
    os.environ["ANTHROPIC_API_KEY"] = anthropic_api_key
    os.environ["SERPER_API_KEY"] = serper_api_key
    os.environ["REPLICATE_API_TOKEN"] = replicate_api_key

    # Create a temporary directory for processing
    with tempfile.TemporaryDirectory() as temp_dir:
        fixed_dir_name = "report_1"
        topic_dir = os.path.join(temp_dir, fixed_dir_name)
        os.makedirs(topic_dir, exist_ok=True)

        from knowledge_storm import (
            STORMWikiRunnerArguments,
            STORMWikiRunner,
            STORMWikiLMConfigs,
        )
        from knowledge_storm.lm import ClaudeModel
        from knowledge_storm.rm import SerperRM

        lm_configs = STORMWikiLMConfigs()
        claude_kwargs = {
            "api_key": os.getenv("ANTHROPIC_API_KEY"),
            "temperature": 1.0,
            "top_p": 0.9,
        }

        conv_simulator_lm = ClaudeModel(
            model="claude-3-haiku-20240307", max_tokens=2096, **claude_kwargs
        )
        question_asker_lm = ClaudeModel(
            model="claude-3-sonnet-20240229", max_tokens=2096, **claude_kwargs
        )
        outline_gen_lm = ClaudeModel(
            model="claude-3-opus-20240229", max_tokens=2096, **claude_kwargs
        )
        article_gen_lm = ClaudeModel(
            model="claude-3-opus-20240229", max_tokens=2096, **claude_kwargs
        )
        article_polish_lm = ClaudeModel(
            model="claude-3-opus-20240229", max_tokens=2096, **claude_kwargs
        )

        lm_configs.set_conv_simulator_lm(conv_simulator_lm)
        lm_configs.set_question_asker_lm(question_asker_lm)
        lm_configs.set_outline_gen_lm(outline_gen_lm)
        lm_configs.set_article_gen_lm(article_gen_lm)
        lm_configs.set_article_polish_lm(article_polish_lm)

        engine_args = STORMWikiRunnerArguments(
            output_dir=topic_dir,
            max_conv_turn=5,
            max_perspective=3,
            search_top_k=10,
            max_thread_num=3,
        )
        rm = SerperRM(serper_search_api_key=os.getenv("SERPER_API_KEY"),
                      query_params={"autocorrect": True, "num": 10, "page": 1})

        runner = STORMWikiRunner(engine_args, lm_configs, rm)
        topic_with_keywords = f"{topic} {keywords}"

        try:
            run_with_retry(runner, topic_with_keywords)
        except Exception as e:
            return None, str(e)

        sanitized_topic = sanitize_filename(topic_with_keywords.replace(" ", "_"))
        generated_dir = os.path.join(topic_dir, sanitized_topic)

        if os.path.exists(generated_dir):
            for file_name in os.listdir(generated_dir):
                src_path = os.path.join(generated_dir, file_name)
                dst_path = os.path.join(topic_dir, file_name)
                shutil.move(src_path, dst_path)
            os.rmdir(generated_dir)

        article_file_path = os.path.join(topic_dir, "storm_gen_article.txt")
        references_file_path = os.path.join(topic_dir, "url_to_info.json")

        if not os.path.exists(article_file_path) or not os.path.exists(references_file_path):
            return None, "Generated files not found"

        with open(article_file_path, 'r') as article_file:
            generated_article = article_file.read()

        with open(references_file_path, 'r') as references_file:
            references_data = json.load(references_file)

        humanized_article = introduce_grammatical_errors(generated_article)

        # Process citations to make them clickable
        processed_article = process_citations(humanized_article)

        formatted_references = format_references(references_data)

        # Build the complete content in memory
        article_content = f"# {topic}\n\n"

        # Add header image
        image_prompt = f"{topic} {keywords}"
        image_url = generate_image(image_prompt, "header_image", replicate_api_key)
        if image_url:
            article_content += f"<img src='{image_url}' alt='{image_prompt}' width='500'/>\n\n"

        # Add the processed article with clickable citations
        article_content += processed_article

        # Add references with an id anchor
        article_content += "\n\n<a id='references'></a>\n## References\n\n"
        article_content += formatted_references

        return article_content, None


def markdown_to_docx(markdown_text, output_filename):
    """Convert markdown text to DOCX format"""


    # Create a new Document
    doc = Document()

    # Add title
    title_match = re.search(r'# (.*?)(\n|$)', markdown_text)
    if title_match:
        title = title_match.group(1).strip()
        doc.add_heading(title, 0)

    # Convert markdown to HTML
    html = markdown.markdown(markdown_text)
    soup = BeautifulSoup(html, 'html.parser')

    # Process content
    for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'img', 'ul', 'ol', 'li']):
        if element.name == 'h1':
            continue  # Skip the title as we've already added it
        elif element.name.startswith('h'):
            level = int(element.name[1])
            doc.add_heading(element.text, level)
        elif element.name == 'p':
            doc.add_paragraph(element.text)
        elif element.name == 'img':
            src = element.get('src')
            if src and src.startswith('http'):
                try:
                    response = requests.get(src)
                    if response.status_code == 200:
                        image_stream = io.BytesIO(response.content)
                        doc.add_picture(image_stream, width=Inches(6))

                        # Add caption
                        alt_text = element.get('alt', '')
                        if alt_text:
                            caption = doc.add_paragraph(alt_text)
                            caption.style = 'Caption'
                except Exception as e:
                    print(f"Error adding image: {e}")
                    doc.add_paragraph(f"[Image: {element.get('alt', 'Image')}]")
        elif element.name == 'ul':
            for li in element.find_all('li', recursive=False):
                doc.add_paragraph(li.text, style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li', recursive=False):
                doc.add_paragraph(li.text, style='List Number')

    # Save the document
    doc.save(output_filename)

    return output_filename


def streamlit_app():
    # Configure the sidebar for API keys
    st.sidebar.title("API Settings")
    st.sidebar.markdown("Enter your API keys")

    anthropic_api_key = st.sidebar.text_input(
        "Anthropic API Key",
        value="",
        type="password"
    )

    serper_api_key = st.sidebar.text_input(
        "Serper API Key",
        value="",
        type="password"
    )

    replicate_api_key = st.sidebar.text_input(
        "Replicate API Token",
        value="",
        type="password"
    )

    # Main app interface
    st.title("Generate Blog Article")
    topic = st.text_input("Enter Blog Title", placeholder="Enter the title of the blog")
    keywords = st.text_input("Enter Keywords", placeholder="Enter keywords related to the blog")

    if st.button("Generate Article"):
        if not topic or not keywords:
            st.warning("Please enter both the blog title and keywords.")
        elif not anthropic_api_key or not serper_api_key or not replicate_api_key:
            st.warning("Please provide all API keys in the sidebar.")
        else:
            with st.spinner("Generating your article... This may take a few minutes"):
                article_content, error = main_in_memory(
                    topic,
                    keywords,
                    anthropic_api_key,
                    serper_api_key,
                    replicate_api_key
                )

                if error:
                    st.error(f"An error occurred: {error}")
                else:
                    st.success(f"Article generated successfully!")
                    st.markdown(article_content, unsafe_allow_html=True)

                    # Optional: Add a download button
                    st.download_button(
                        label="Download Article as Markdown",
                        data=article_content,
                        file_name=f"{topic.replace(' ', '_')}.md",
                        mime="text/markdown"
                    )


if __name__ == "__main__":
    streamlit_app()


